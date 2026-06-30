import hashlib
import importlib.util
import json
import sys
from pathlib import Path
from types import ModuleType

from app.schemas.uir import UIRDocument

ROOT = Path(__file__).resolve().parents[2]
SCRIPTS_DIR = ROOT / "scripts"


def load_script(name: str) -> ModuleType:
    path = SCRIPTS_DIR / f"{name}.py"
    if str(SCRIPTS_DIR) not in sys.path:
        sys.path.insert(0, str(SCRIPTS_DIR))
    spec = importlib.util.spec_from_file_location(name, path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"cannot load script: {path}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[name] = module
    spec.loader.exec_module(module)
    return module


def test_build_uir_matches_existing_schema() -> None:
    common = load_script("real_world_uir_common")
    source_bytes = "官方公开内容".encode()
    uir = common.build_uir(
        source={
            "source_id": "real_policy_001_public_policy",
            "doc_type": "policy_doc",
            "source_url": "https://example.gov.cn/policy.html",
            "source_site": "example.gov.cn",
        },
        title="公开政策",
        blocks=[
            {"type": "heading", "level": 1, "text": "公开政策"},
            {"type": "paragraph", "text": "发布日期：2026-06-29"},
            {"type": "paragraph", "text": "本政策自发布之日起施行。"},
        ],
        source_bytes=source_bytes,
        retrieved_at="2026-06-29T00:00:00+08:00",
        source_format="html",
        extraction_method="beautifulsoup_lxml",
    )

    parsed = UIRDocument.model_validate(uir)
    assert parsed.doc_id == "real_policy_001_public_policy"
    assert [block.block_id for block in parsed.blocks] == [
        "real_policy_001_public_policy_b001",
        "real_policy_001_public_policy_b002",
        "real_policy_001_public_policy_b003",
    ]
    assert uir["metadata"]["doc_type"] == "policy_doc"
    assert uir["metadata"]["domain"] == "policy_doc"
    assert uir["metadata"]["extraction_version"] == "0.1.0"
    assert uir["metadata"]["source_sha256"] == hashlib.sha256(source_bytes).hexdigest()
    assert uir["source"] == {
        "source_type": "real_world_public_document",
        "source_name": "real_policy_001_public_policy",
        "upstream_agents": ["real_world_uir_builder"],
    }


def test_build_uir_caps_very_long_documents_and_records_truncation() -> None:
    common = load_script("real_world_uir_common")
    blocks = [
        {"type": "paragraph", "text": f"第 {index} 行"}
        for index in range(common.MAX_UIR_BLOCKS + 10)
    ]

    uir = common.build_uir(
        source={
            "source_id": "real_policy_long",
            "doc_type": "policy_doc",
            "source_url": "https://example.gov.cn/long.pdf",
            "source_site": "example.gov.cn",
        },
        title="长文档",
        blocks=blocks,
        source_bytes=b"long",
        retrieved_at="2026-06-29T00:00:00+08:00",
        source_format="pdf",
        extraction_method="pymupdf_text",
    )

    assert len(uir["blocks"]) == common.MAX_UIR_BLOCKS
    assert uir["metadata"]["extraction_truncated"] is True
    assert uir["metadata"]["extracted_block_count"] == common.MAX_UIR_BLOCKS + 10


def test_html_extractor_emits_content_blocks_in_document_order() -> None:
    extractor = load_script("extract_html_to_uir")
    html = """
    <html>
      <head><title>站点标题</title><script>secret()</script></head>
      <body>
        <nav>栏目导航</nav>
        <article>
          <h1>示例政策</h1>
          <p>发布日期：2026年6月29日</p>
          <ul><li>第一项</li><li>第二项</li></ul>
          <table>
            <tr><th>项目</th><th>金额</th></tr>
            <tr><td>A</td><td>100</td></tr>
          </table>
        </article>
        <footer>备案号</footer>
      </body>
    </html>
    """.encode()

    result = extractor.extract_html(html, source_url="https://example.gov.cn/doc")

    assert result.title == "示例政策"
    assert [block["type"] for block in result.blocks] == [
        "heading",
        "paragraph",
        "list",
        "table",
    ]
    assert result.blocks[2]["attributes"]["items"] == ["第一项", "第二项"]
    assert result.blocks[3]["attributes"]["rows"] == [
        {"field": "项目", "value": "金额"},
        {"field": "A", "value": "100"},
    ]
    assert result.metadata["发布日期"] == "2026年6月29日"
    all_text = str(result.blocks)
    assert "栏目导航" not in all_text
    assert "备案号" not in all_text
    assert "secret" not in all_text


def test_html_extractor_rejects_empty_content() -> None:
    extractor = load_script("extract_html_to_uir")

    result = extractor.extract_html(
        b"<html><body><nav>navigation only</nav></body></html>",
        source_url="https://example.gov.cn/empty",
    )

    assert result.status == "rejected"
    assert result.reason == "empty_extraction"
    assert result.blocks == []


def test_html_extractor_treats_presentation_table_as_layout() -> None:
    extractor = load_script("extract_html_to_uir")
    html = """
    <article>
      <h1>布局表格中的政策</h1>
      <table role="presentation">
        <tr><td>
          <p>第一条 正文内容。</p>
          <p>第二条 继续说明。</p>
        </td></tr>
      </table>
    </article>
    """.encode()

    result = extractor.extract_html(html, source_url="https://example.gov.cn/layout")

    assert [block["type"] for block in result.blocks] == [
        "heading",
        "paragraph",
        "paragraph",
    ]
    assert result.blocks[-1]["text"] == "第二条 继续说明。"


def test_html_extractor_prefers_known_government_content_container() -> None:
    extractor = load_script("extract_html_to_uir")
    html = """
    <html><body>
      <p>页面栏目说明，不属于正文。</p>
      <div id="ivs_content" class="xxgk_content_nr">
        <h1>项目申报指南</h1>
        <p>一、申报范围</p>
        <p>二、申报条件</p>
      </div>
    </body></html>
    """.encode()

    result = extractor.extract_html(html, source_url="https://example.gov.cn/guide")

    assert [block["text"] for block in result.blocks] == [
        "项目申报指南",
        "一、申报范围",
        "二、申报条件",
    ]


def test_html_extractor_falls_back_to_text_nodes_for_nonsemantic_article() -> None:
    extractor = load_script("extract_html_to_uir")
    html = """
    <article>
      <h1>设备更新通知</h1>
      <div>第一条 推进重点领域设备更新。<br>第二条 实施贷款贴息。</div>
    </article>
    """.encode()

    result = extractor.extract_html(html, source_url="https://example.gov.cn/nonsemantic")

    assert [block["text"] for block in result.blocks] == [
        "设备更新通知",
        "第一条 推进重点领域设备更新。",
        "第二条 实施贷款贴息。",
    ]


def test_pdf_extractor_preserves_page_anchors(tmp_path: Path) -> None:
    import fitz

    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Public Policy", fontsize=18)
    page.insert_text((72, 110), "Published: 2026-06-29", fontsize=11)
    page.insert_text(
        (72, 145),
        "This policy defines a reproducible evaluation process.",
        fontsize=11,
    )
    pdf_path = tmp_path / "text.pdf"
    document.save(pdf_path)
    document.close()
    extractor = load_script("extract_pdf_to_uir")

    result = extractor.extract_pdf(pdf_path.read_bytes())

    assert result.status == "extracted"
    assert result.title == "Public Policy"
    assert len(result.blocks) >= 3
    assert all(block["source_anchor"]["page"] == 1 for block in result.blocks)


def test_pdf_extractor_skips_scanned_or_image_only_pdf(tmp_path: Path) -> None:
    import fitz

    document = fitz.open()
    document.new_page()
    pdf_path = tmp_path / "image-only.pdf"
    document.save(pdf_path)
    document.close()
    extractor = load_script("extract_pdf_to_uir")

    result = extractor.extract_pdf(pdf_path.read_bytes())

    assert result.status == "skipped"
    assert result.reason == "unsupported_scanned_pdf"
    assert result.blocks == []


def test_docx_extractor_maps_headings_paragraphs_and_tables(tmp_path: Path) -> None:
    from docx import Document

    document = Document()
    document.add_heading("Service Guide", level=1)
    document.add_paragraph("Published: 2026-06-29")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Window"
    table.cell(1, 1).text = "Online"
    docx_path = tmp_path / "guide.docx"
    document.save(docx_path)
    extractor = load_script("extract_docx_to_uir")

    result = extractor.extract_docx(docx_path.read_bytes())

    assert result.title == "Service Guide"
    assert [block["type"] for block in result.blocks] == [
        "heading",
        "paragraph",
        "table",
    ]
    assert result.blocks[-1]["attributes"]["rows"][-1] == {
        "field": "Window",
        "value": "Online",
    }


class FakeResponse:
    def __init__(
        self,
        content: bytes,
        *,
        status_code: int = 200,
        content_type: str = "text/html; charset=utf-8",
    ) -> None:
        self.content = content
        self.status_code = status_code
        self.headers = {
            "Content-Type": content_type,
            "Content-Length": str(len(content)),
        }

    def iter_content(self, chunk_size: int = 65536):
        for index in range(0, len(self.content), chunk_size):
            yield self.content[index : index + chunk_size]


class FakeSession:
    def __init__(self, responses: dict[str, FakeResponse]) -> None:
        self.responses = responses

    def get(self, url: str, **_: object) -> FakeResponse:
        return self.responses[url]


def write_manifest(path: Path, items: list[dict[str, object]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(
            {
                "dataset_version": "0.1.0",
                "created_at": "2026-06-29T00:00:00+08:00",
                "description": "test",
                "items": items,
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )


def test_real_world_dataset_has_at_least_thirty_official_sources() -> None:
    manifest_path = ROOT / "examples" / "real_world" / "sources" / "source_manifest.json"

    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    items = manifest["items"]

    assert len(items) >= 30
    source_ids = [item["source_id"] for item in items]
    source_urls = [item["source_url"] for item in items]
    uir_paths = [item["uir_path"] for item in items]
    assert len(source_ids) == len(set(source_ids))
    assert len(source_urls) == len(set(source_urls))
    assert len(uir_paths) == len(set(uir_paths))
    for item in items:
        assert item["source_url"].startswith("https://"), item["source_id"]
        assert item["license_note"], item["source_id"]
        assert item["source_site"], item["source_id"]


def test_collector_caches_content_and_updates_manifest(tmp_path: Path) -> None:
    collector = load_script("collect_real_world_sources")
    manifest_path = tmp_path / "sources" / "source_manifest.json"
    cache_dir = tmp_path / "raw_cache"
    content = "<html><article><h1>公开政策</h1></article></html>".encode()
    url = "https://example.gov.cn/policy.html"
    write_manifest(
        manifest_path,
        [
            {
                "source_id": "real_policy_001",
                "doc_type": "policy_doc",
                "title": "公开政策",
                "source_url": url,
                "source_site": "example.gov.cn",
                "source_format": "html",
                "retrieval_method": "requests",
                "status": "planned",
                "license_note": "official public webpage",
            }
        ],
    )

    summary = collector.collect_manifest(
        manifest_path=manifest_path,
        cache_dir=cache_dir,
        session=FakeSession({url: FakeResponse(content)}),
        max_bytes=1024,
        retrieved_at_factory=lambda: "2026-06-29T10:00:00+08:00",
    )

    saved_manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    item = saved_manifest["items"][0]
    assert summary == {"fetched": 1, "skipped": 0, "failed": 0}
    assert item["status"] == "fetched"
    assert item["source_sha256"] == hashlib.sha256(content).hexdigest()
    assert item["retrieved_at"] == "2026-06-29T10:00:00+08:00"
    assert (cache_dir / item["cached_path"]).read_bytes() == content


def test_collector_rejects_unsafe_url_and_oversized_response(tmp_path: Path) -> None:
    collector = load_script("collect_real_world_sources")
    manifest_path = tmp_path / "source_manifest.json"
    cache_dir = tmp_path / "raw_cache"
    large_url = "https://example.gov.cn/large.pdf"
    write_manifest(
        manifest_path,
        [
            {
                "source_id": "bad_scheme",
                "doc_type": "general_doc",
                "title": "bad",
                "source_url": "file:///etc/passwd",
                "source_site": "local",
                "source_format": "html",
                "retrieval_method": "requests",
                "status": "planned",
                "license_note": "none",
            },
            {
                "source_id": "too_large",
                "doc_type": "policy_doc",
                "title": "large",
                "source_url": large_url,
                "source_site": "example.gov.cn",
                "source_format": "pdf",
                "retrieval_method": "requests",
                "status": "planned",
                "license_note": "official",
            },
        ],
    )

    summary = collector.collect_manifest(
        manifest_path=manifest_path,
        cache_dir=cache_dir,
        session=FakeSession(
            {large_url: FakeResponse(b"x" * 20, content_type="application/pdf")}
        ),
        max_bytes=10,
    )

    items = json.loads(manifest_path.read_text(encoding="utf-8"))["items"]
    assert summary == {"fetched": 0, "skipped": 0, "failed": 2}
    assert items[0]["failure_reason"] == "invalid_source_url"
    assert items[1]["failure_reason"] == "content_too_large"
    assert list(cache_dir.glob("*")) == []


def test_collector_refetches_extracted_source_when_cache_was_removed(
    tmp_path: Path,
) -> None:
    collector = load_script("collect_real_world_sources")
    manifest_path = tmp_path / "source_manifest.json"
    cache_dir = tmp_path / "raw_cache"
    url = "https://example.gov.cn/rebuild.html"
    content = b"<html><article><h1>Rebuilt</h1></article></html>"
    write_manifest(
        manifest_path,
        [
            {
                "source_id": "real_general_rebuild",
                "doc_type": "general_doc",
                "title": "Rebuilt",
                "source_url": url,
                "source_site": "example.gov.cn",
                "source_format": "html",
                "retrieval_method": "requests",
                "status": "extracted",
                "license_note": "official",
                "cached_path": "real_general_rebuild.html",
            }
        ],
    )

    summary = collector.collect_manifest(
        manifest_path=manifest_path,
        cache_dir=cache_dir,
        session=FakeSession({url: FakeResponse(content)}),
    )

    assert summary["fetched"] == 1
    assert (cache_dir / "real_general_rebuild.html").read_bytes() == content


def test_collector_rejects_content_type_mismatch(tmp_path: Path) -> None:
    collector = load_script("collect_real_world_sources")
    manifest_path = tmp_path / "source_manifest.json"
    cache_dir = tmp_path / "raw_cache"
    url = "https://example.gov.cn/not-really.pdf"
    write_manifest(
        manifest_path,
        [
            {
                "source_id": "real_policy_wrong_type",
                "doc_type": "policy_doc",
                "title": "Wrong type",
                "source_url": url,
                "source_site": "example.gov.cn",
                "source_format": "pdf",
                "retrieval_method": "requests",
                "status": "planned",
                "license_note": "official",
            }
        ],
    )

    collector.collect_manifest(
        manifest_path=manifest_path,
        cache_dir=cache_dir,
        session=FakeSession({url: FakeResponse(b"<html>error</html>")}),
    )

    item = json.loads(manifest_path.read_text(encoding="utf-8"))["items"][0]
    assert item["status"] == "failed"
    assert item["failure_reason"] == "content_type_mismatch"


def test_build_dataset_dispatches_html_and_pdf_and_writes_reports(tmp_path: Path) -> None:
    import fitz

    builder = load_script("build_real_world_uir")
    manifest_path = tmp_path / "sources" / "source_manifest.json"
    cache_dir = tmp_path / "raw_cache"
    uir_dir = tmp_path / "uir"
    reports_dir = tmp_path / "reports"
    cache_dir.mkdir(parents=True)

    html_content = """
    <article>
      <h1>公开服务指南</h1>
      <p>发布日期：2026年6月29日</p>
      <p>第一章 申请条件</p>
    </article>
    """.encode()
    (cache_dir / "real_general_001.html").write_bytes(html_content)

    pdf_document = fitz.open()
    page = pdf_document.new_page()
    page.insert_text((72, 72), "Procurement Notice", fontsize=18)
    page.insert_text((72, 110), "Project number: P-001", fontsize=11)
    page.insert_text((72, 145), "Budget amount: 100000 yuan", fontsize=11)
    pdf_path = cache_dir / "real_procurement_001.pdf"
    pdf_document.save(pdf_path)
    pdf_document.close()
    pdf_content = pdf_path.read_bytes()

    write_manifest(
        manifest_path,
        [
            {
                "source_id": "real_general_001",
                "doc_type": "general_doc",
                "title": "公开服务指南",
                "source_url": "https://example.gov.cn/guide.html",
                "source_site": "example.gov.cn",
                "source_format": "html",
                "retrieval_method": "requests",
                "status": "fetched",
                "license_note": "official",
                "cached_path": "real_general_001.html",
                "retrieved_at": "2026-06-29T10:00:00+08:00",
                "source_sha256": hashlib.sha256(html_content).hexdigest(),
            },
            {
                "source_id": "real_procurement_001",
                "doc_type": "procurement_doc",
                "title": "Procurement Notice",
                "source_url": "https://example.gov.cn/bid.pdf",
                "source_site": "example.gov.cn",
                "source_format": "pdf",
                "retrieval_method": "requests",
                "status": "fetched",
                "license_note": "official",
                "cached_path": "real_procurement_001.pdf",
                "retrieved_at": "2026-06-29T10:00:00+08:00",
                "source_sha256": hashlib.sha256(pdf_content).hexdigest(),
            },
        ],
    )

    report = builder.build_dataset(
        manifest_path=manifest_path,
        cache_dir=cache_dir,
        uir_dir=uir_dir,
        reports_dir=reports_dir,
    )

    assert report["totals"] == {
        "sources": 2,
        "extracted": 2,
        "rejected": 0,
        "skipped": 0,
    }
    assert report["by_format"] == {"html": 1, "pdf": 1}
    assert report["collection_totals"] == {
        "downloaded": 2,
        "skipped": 0,
        "failed": 0,
    }
    general = json.loads(
        (uir_dir / "general" / "real_general_001.json").read_text(encoding="utf-8")
    )
    procurement = json.loads(
        (uir_dir / "procurement" / "real_procurement_001.json").read_text(encoding="utf-8")
    )
    assert UIRDocument.model_validate(general).metadata["source_format"] == "html"
    assert UIRDocument.model_validate(procurement).metadata["source_format"] == "pdf"
    assert (reports_dir / "extraction_report.json").is_file()
    assert "Extracted | 2" in (reports_dir / "extraction_report.md").read_text(
        encoding="utf-8"
    )
    statuses = [
        item["status"]
        for item in json.loads(manifest_path.read_text(encoding="utf-8"))["items"]
    ]
    assert statuses == ["extracted", "extracted"]


def make_valid_uir() -> dict[str, object]:
    common = load_script("real_world_uir_common")
    return common.build_uir(
        source={
            "source_id": "real_policy_001",
            "doc_type": "policy_doc",
            "source_url": "https://example.gov.cn/policy.html",
            "source_site": "example.gov.cn",
        },
        title="公开政策",
        blocks=[
            {"type": "heading", "level": 1, "text": "公开政策"},
            {"type": "paragraph", "text": "发布日期：2026-06-29"},
            {"type": "paragraph", "text": "第一条 本办法规定公开事项。"},
        ],
        source_bytes=b"official public content",
        retrieved_at="2026-06-29T10:00:00+08:00",
        source_format="html",
        extraction_method="beautifulsoup_lxml",
    )


def test_validator_accepts_traceable_existing_schema_uir() -> None:
    validator = load_script("validate_real_world_uir")

    findings = validator.validate_uir_data(make_valid_uir())

    assert findings == []


def test_validator_flags_sensitive_data_and_unreviewed_low_confidence_hint() -> None:
    validator = load_script("validate_real_world_uir")
    uir = make_valid_uir()
    uir["blocks"][2]["text"] += " 联系人手机：13800138000"
    uir["metadata"]["hints"] = {
        "candidate_fields": [
            {
                "field": "issuer",
                "value": "某机构",
                "confidence": 0.4,
                "evidence_text": "某机构",
                "evidence_block_ids": ["real_policy_001_b003"],
                "review_required": False,
            }
        ]
    }

    findings = validator.validate_uir_data(uir)
    codes = {finding["code"] for finding in findings}

    assert "possible_personal_sensitive_information" in codes
    assert "low_confidence_without_review" in codes


def test_sensitive_scanner_ignores_numeric_pdf_coordinates() -> None:
    validator = load_script("validate_real_world_uir")
    uir = make_valid_uir()
    uir["blocks"][0]["source_anchor"] = {
        "page": 1,
        "bbox": [90.0, 323.18896484375, 505.18707275390625, 339.5058288574219],
    }

    findings = validator.scan_sensitive_information(uir)

    assert findings == []


def test_sensitive_scanner_detects_email_adjacent_to_cjk_text() -> None:
    validator = load_script("validate_real_world_uir")
    uir = make_valid_uir()
    uir["blocks"][2]["text"] += " 反馈至jiayue@miitec.cn"

    findings = validator.scan_sensitive_information(uir)

    assert any(
        finding["code"] == "possible_personal_sensitive_information"
        and finding["message"] == "matched personal_email"
        for finding in findings
    )


def test_validator_rejects_bad_hash_empty_block_and_mojibake() -> None:
    validator = load_script("validate_real_world_uir")
    uir = make_valid_uir()
    uir["metadata"]["source_sha256"] = "not-a-hash"
    uir["blocks"][1]["text"] = ""
    uir["blocks"][2]["text"] = "æ–‡æ¡£å†…å®¹"

    findings = validator.validate_uir_data(uir)
    codes = {finding["code"] for finding in findings}

    assert {"invalid_source_sha256", "empty_block", "possible_mojibake"} <= codes


def test_validation_dataset_moves_rejected_files_inside_dataset_root(tmp_path: Path) -> None:
    validator = load_script("validate_real_world_uir")
    uir_root = tmp_path / "uir"
    policy_dir = uir_root / "policy"
    reports_dir = tmp_path / "reports"
    policy_dir.mkdir(parents=True)
    invalid_path = policy_dir / "invalid.json"
    invalid_path.write_text(
        json.dumps({"uir_version": "1.0", "doc_id": "invalid"}),
        encoding="utf-8",
    )

    report = validator.validate_dataset(
        uir_dir=uir_root,
        reports_dir=reports_dir,
        move_rejected=True,
    )

    assert report["totals"]["failed"] == 1
    assert report["totals"]["missing_fields"] > 0
    assert report["totals"]["empty_or_mojibake"] > 0
    assert not invalid_path.exists()
    assert (uir_root / "_rejected" / "invalid.json").is_file()
    assert (reports_dir / "validation_report.json").is_file()


def test_evaluator_routes_procurement_to_dedicated_catalog_and_continues_after_failure(
    tmp_path: Path,
) -> None:
    import httpx

    evaluator = load_script("eval_real_world_uir")
    uir_dir = tmp_path / "uir"
    (uir_dir / "policy").mkdir(parents=True)
    (uir_dir / "procurement").mkdir(parents=True)
    failed_uir = make_valid_uir()
    failed_uir["doc_id"] = "real_policy_failed"
    failed_uir["metadata"]["doc_type"] = "policy_doc"
    failed_uir["metadata"]["domain"] = "policy_doc"
    good_uir = make_valid_uir()
    good_uir["doc_id"] = "real_procurement_001"
    good_uir["metadata"]["doc_type"] = "procurement_doc"
    good_uir["metadata"]["domain"] = "procurement_doc"
    (uir_dir / "policy" / "real_policy_failed.json").write_text(
        json.dumps(failed_uir, ensure_ascii=False),
        encoding="utf-8",
    )
    (uir_dir / "procurement" / "real_procurement_001.json").write_text(
        json.dumps(good_uir, ensure_ascii=False),
        encoding="utf-8",
    )
    task_payloads: list[dict[str, object]] = []

    def handler(request: httpx.Request) -> httpx.Response:
        assert request.headers["X-API-Key"] == "test-secret"
        path = request.url.path
        if path == "/api/v1/documents/import":
            payload = json.loads(request.content)
            if payload["uir"]["doc_id"] == "real_policy_failed":
                return httpx.Response(400, json={"detail": "invalid pilot"})
            return httpx.Response(
                200,
                json={"doc_id": payload["uir"]["doc_id"], "status": "imported", "block_count": 3},
            )
        if path == "/api/v1/tasks":
            payload = json.loads(request.content)
            task_payloads.append(payload)
            return httpx.Response(200, json={"task_id": "task-1", "status": "created"})
        if path == "/api/v1/tasks/task-1/execute":
            return httpx.Response(
                200,
                json={
                    "task_id": "task-1",
                    "status": "review_required",
                    "report_paths": {},
                    "package_zip_path": "tasks/task-1/package.zip",
                    "review_required_count": 1,
                    "unmapped_required_count": 0,
                },
            )
        if path.endswith("/reports/mapping"):
            return httpx.Response(
                200,
                json={
                    "task_id": "task-1",
                    "schema_id": "procurement_doc",
                    "summary": {},
                    "mappings": [{"risk_level": "high"}],
                    "unmapped": [],
                    "review_required_items": [{"field": "issuer"}],
                },
            )
        if path.endswith("/reports/validation"):
            return httpx.Response(
                200,
                json={
                    "task_id": "task-1",
                    "schema_id": "procurement_doc",
                    "passed": True,
                    "summary": {},
                    "issues": [],
                },
            )
        if path.endswith("/reports/content-organization"):
            return httpx.Response(200, json={"task_id": "task-1", "chunk_count": 1})
        if path.endswith("/reports/chunks"):
            return httpx.Response(200, json={"items": [{"chunk_id": "c1"}], "total": 1})
        if path.endswith("/reports/verifier"):
            return httpx.Response(
                200,
                json={
                    "task_id": "task-1",
                    "passed": True,
                    "checks": [],
                    "errors": [],
                    "warnings": [],
                },
            )
        if path == "/api/v1/tasks/task-1/package":
            return httpx.Response(200, json={"package_id": "pkg-1"})
        if path == "/api/v1/tasks/task-1/package/download":
            return httpx.Response(200, content=b"PK-test-package")
        raise AssertionError(f"unexpected request: {request.method} {path}")

    client = httpx.Client(
        base_url="http://testserver",
        headers={"X-API-Key": "test-secret"},
        transport=httpx.MockTransport(handler),
    )
    reports_dir = tmp_path / "reports"
    packages_dir = tmp_path / "packages"

    report = evaluator.evaluate_dataset(
        uir_dir=uir_dir,
        reports_dir=reports_dir,
        packages_dir=packages_dir,
        client=client,
    )

    assert report["dataset_size"] == 2
    assert report["import_pass_count"] == 1
    assert report["task_execute_pass_count"] == 1
    assert report["package_verify_pass_count"] == 1
    assert report["mapping_review_required_count"] == 1
    assert report["high_risk_mapping_count"] == 1
    assert report["by_doc_type_metrics"]["procurement_doc"] == {
        "dataset_size": 1,
        "import_pass_count": 1,
        "task_execute_pass_count": 1,
        "package_verify_pass_count": 1,
        "validation_pass_count": 1,
    }
    assert report["typical_success_cases"] == ["real_procurement_001"]
    assert report["typical_failure_cases"][0]["doc_id"] == "real_policy_failed"
    assert task_payloads[0]["schema_id"] == "procurement_doc"
    assert task_payloads[0]["schema_version"] == "1.0.0"
    assert task_payloads[0]["template_id"] == "procurement_doc_base_v1"
    assert task_payloads[0]["template_version"] == "1.0.0"
    procurement_item = next(
        item for item in report["items"] if item["doc_type"] == "procurement_doc"
    )
    assert procurement_item["catalog_status"] == "available"
    assert (packages_dir / "real_procurement_001.zip").read_bytes() == b"PK-test-package"
    rendered_report = (reports_dir / "real_world_eval_report.json").read_text(
        encoding="utf-8"
    )
    assert "test-secret" not in rendered_report
    assert "invalid pilot" in rendered_report


def test_evaluator_reports_missing_procurement_catalog_without_fallback(
    tmp_path: Path,
    monkeypatch,
) -> None:
    import httpx

    evaluator = load_script("eval_real_world_uir")
    uir_dir = tmp_path / "uir"
    (uir_dir / "procurement").mkdir(parents=True)
    uir = make_valid_uir()
    uir["doc_id"] = "real_procurement_missing_catalog"
    uir["metadata"]["doc_type"] = "procurement_doc"
    uir["metadata"]["domain"] = "procurement_doc"
    (uir_dir / "procurement" / "real_procurement_missing_catalog.json").write_text(
        json.dumps(uir, ensure_ascii=False),
        encoding="utf-8",
    )
    missing_schema = tmp_path / "missing_procurement_schema.json"
    monkeypatch.setitem(
        evaluator.DOCUMENT_CATALOG["procurement_doc"],
        "schema_path",
        missing_schema,
    )

    def handler(request: httpx.Request) -> httpx.Response:
        raise AssertionError(f"catalog failure must happen before HTTP: {request.url.path}")

    with httpx.Client(
        base_url="http://testserver",
        transport=httpx.MockTransport(handler),
    ) as client:
        report = evaluator.evaluate_dataset(
            uir_dir=uir_dir,
            reports_dir=tmp_path / "reports",
            packages_dir=tmp_path / "packages",
            client=client,
        )

    assert report["import_pass_count"] == 0
    assert report["task_execute_pass_count"] == 0
    item = report["items"][0]
    assert item["catalog_status"] == "missing_fixture"
    assert "procurement_doc" in item["error"]
    assert "general_doc" not in item["error"]
